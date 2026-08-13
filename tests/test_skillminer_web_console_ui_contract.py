from __future__ import annotations

import base64
import hashlib
import io
import json
import sys
import threading
from pathlib import Path

import pytest
import yaml

SKILLMINER_ROOT = Path(__file__).resolve().parents[1] / "teamEvolver" / "skillminer"
if str(SKILLMINER_ROOT) not in sys.path:
    sys.path.insert(0, str(SKILLMINER_ROOT))

from web_console import server  # noqa: E402


def test_mining_model_settings_can_be_saved_without_returning_secret(tmp_path, monkeypatch):
    config_path = tmp_path / ".hermes_home" / "config.yaml"
    config_path.parent.mkdir(parents=True)
    config_path.write_text(
        "model:\n  default: old-model\n  provider: custom\n  base_url: https://old.example/v1\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(server.rp, "resolve_ark_key", lambda: "")

    saved = server.save_mining_model_settings({
        "model": "new-model",
        "base_url": "https://model.example/v1/",
        "max_tokens": 4096,
        "temperature": 0.3,
        "api_key": "test-secret-key",
    })

    assert saved == {
        "provider": "openai-compatible",
        "id": "new-model",
        "model": "new-model",
        "base_url": "https://model.example/v1",
        "max_tokens": 4096,
        "temperature": 0.3,
        "api_key_present": True,
        "configured": True,
    }
    assert "api_key" not in saved
    persisted = yaml.safe_load(config_path.read_text(encoding="utf-8"))
    assert persisted["model"]["default"] == "new-model"
    assert persisted["model"]["provider"] == "custom"
    assert persisted["model"]["api_mode"] == "chat_completions"
    assert persisted["model"]["api_key"] == "test-secret-key"
    assert config_path.stat().st_mode & 0o777 == 0o600


def test_saving_mining_model_settings_drops_legacy_evolution_hooks(tmp_path, monkeypatch):
    config_path = tmp_path / ".hermes_home" / "config.yaml"
    config_path.parent.mkdir(parents=True)
    config_path.write_text(
        """
model:
  default: old-model
  base_url: https://old.example/v1
hooks:
  on_session_end:
    - command: python push_session.py
skills:
  external_dirs:
    - /Users/example/.hermes/team_skills/skillgene
""".lstrip(),
        encoding="utf-8",
    )
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(server.rp, "resolve_ark_key", lambda: "")

    server.save_mining_model_settings({
        "model": "new-model",
        "base_url": "https://model.example/v1",
        "max_tokens": 4096,
        "temperature": 0.2,
    })

    persisted = yaml.safe_load(config_path.read_text(encoding="utf-8"))
    assert "hooks" not in persisted
    assert persisted["hooks_auto_accept"] is False
    assert "skills" not in persisted


def test_mining_model_test_uses_current_form_values(tmp_path, monkeypatch):
    config_path = tmp_path / ".hermes_home" / "config.yaml"
    config_path.parent.mkdir(parents=True)
    config_path.write_text(
        "model:\n  default: saved-model\n  provider: custom\n  base_url: https://saved.example/v1\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(server.rp, "resolve_ark_key", lambda: "")
    captured = {}

    class FakeResponse:
        def __enter__(self):
            return self

        def __exit__(self, *_args):
            return False

        def read(self, _limit):
            return json.dumps({"choices": [{"message": {"content": "OK"}}]}).encode()

    def fake_urlopen(request, timeout):
        captured["url"] = request.full_url
        captured["authorization"] = request.get_header("Authorization")
        captured["payload"] = json.loads(request.data.decode("utf-8"))
        captured["timeout"] = timeout
        return FakeResponse()

    monkeypatch.setattr(server.urllib.request, "urlopen", fake_urlopen)
    result = server.test_mining_model_settings({
        "model": "edited-model",
        "base_url": "https://edited.example/v1",
        "api_key": "edited-secret-key",
        "max_tokens": 8192,
        "temperature": 0.1,
    })

    assert result["ok"] is True
    assert result["model"] == "edited-model"
    assert result["response"] == "OK"
    assert captured["url"] == "https://edited.example/v1/chat/completions"
    assert captured["authorization"] == "Bearer edited-secret-key"
    assert captured["payload"]["model"] == "edited-model"
    assert captured["timeout"] == 60


def test_config_schema_reports_real_source_and_artifact_readiness(tmp_path, monkeypatch):
    input_dir = tmp_path / "data" / "input"
    input_dir.mkdir(parents=True)
    (input_dir / ".gitkeep").write_text("", encoding="utf-8")
    (input_dir / "guide.md").write_text("usable knowledge", encoding="utf-8")

    skill_dir = tmp_path / "compiled_skill" / "support-skill"
    skill_dir.mkdir(parents=True)
    (skill_dir / "SKILL.md").write_text("# Skill", encoding="utf-8")
    (skill_dir / "EVALUATION.md").write_text("# Evaluation", encoding="utf-8")
    (skill_dir / "benchmark.jsonl").write_text('{"id":"q1"}\n{"id":"q2"}\n', encoding="utf-8")

    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    monkeypatch.setattr(server.li, "PROJECT_ROOT", tmp_path)

    schema = server.build_config_schema()

    assert schema["default_input_dir"] == "data/input"
    assert schema["input_sources"] == [{
        "path": "data/input",
        "document_count": 1,
        "total_bytes": len("usable knowledge"),
        "ready": True,
        "ingestion": {
            "schema_version": 1,
            "source_path": "data/input",
            "batch_id": "",
            "status": "ready",
            "stage": "complete",
            "progress": 100,
            "processed_files": 0,
            "total_files": 0,
            "current_file": "",
            "error": "",
            "started_at": "",
            "updated_at": "",
            "finished_at": "",
        },
    }]
    assert schema["compiled_skill_details"] == [{
        "name": "support-skill",
        "has_skill": True,
        "has_evaluation": True,
        "has_benchmark": True,
        "question_count": 2,
    }]


def test_mining_history_merges_rounds_with_next_run_preexisting_snapshot(tmp_path, monkeypatch):
    round_skill = (
        tmp_path / "reflection_rounds" / "20260805_150618_777997" /
        "round_1" / "compiled_skill" / "support-skill"
    )
    round_skill.mkdir(parents=True)
    (round_skill / "SKILL.md").write_text("# Skill", encoding="utf-8")
    (round_skill / "EVALUATION.md").write_text("# Evaluation", encoding="utf-8")

    final_skill = (
        tmp_path / "run_history" / "20260809_223901_916214" / "preexisting" /
        "compiled_skill" / "support-skill"
    )
    final_skill.mkdir(parents=True)
    (final_skill / "SKILL.md").write_text("# Final Skill", encoding="utf-8")
    (final_skill / "EVALUATION.md").write_text("# Final Evaluation", encoding="utf-8")
    (final_skill / "benchmark.jsonl").write_text('{"id":"q1"}\n{"id":"q2"}\n', encoding="utf-8")
    report = (
        tmp_path / "run_history" / "20260809_223901_916214" / "preexisting" /
        "semantic_reports" / "report.md"
    )
    report.parent.mkdir(parents=True)
    report.write_text("# Report", encoding="utf-8")

    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    runs = server.list_mining_runs()

    assert len(runs) == 1
    assert runs[0]["run_id"] == "20260805_150618_777997"
    assert runs[0]["started_at"] == "2026-08-05T15:06:18"
    assert runs[0]["rounds"][0]["round"] == 1
    assert runs[0]["skills"] == [{
        "name": "support-skill",
        "has_skill": True,
        "has_evaluation": True,
        "has_benchmark": True,
        "question_count": 2,
    }]
    assert any(item["kind"] == "semantic" for item in runs[0]["final_artifacts"])


def test_history_artifact_reader_is_path_safe(tmp_path, monkeypatch):
    artifact = tmp_path / "reflection_rounds" / "run-1" / "round_1" / "compiled_skill" / "x" / "SKILL.md"
    artifact.parent.mkdir(parents=True)
    artifact.write_text("# Safe", encoding="utf-8")
    outside = tmp_path / "secret.txt"
    outside.write_text("secret", encoding="utf-8")
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    result = server.read_history_artifact(
        "reflection_rounds/run-1/round_1/compiled_skill/x/SKILL.md"
    )
    assert result["content"] == "# Safe"

    try:
        server.read_history_artifact("secret.txt")
    except ValueError:
        pass
    else:
        raise AssertionError("artifact reader allowed a path outside history roots")


def test_empty_input_is_rejected_before_worker_start(tmp_path, monkeypatch):
    input_dir = tmp_path / "data" / "empty"
    input_dir.mkdir(parents=True)
    (input_dir / ".gitkeep").write_text("", encoding="utf-8")
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    manager = server.RunManager()
    ok, message = manager.start({"input_dir": "data/empty"})

    assert ok is False
    assert "没有可用于挖掘的文档" in message
    assert manager.thread is None
    assert manager.state == "idle"


def test_compiled_skill_selection_is_exact_and_path_safe(tmp_path, monkeypatch):
    skill_dir = tmp_path / "compiled_skill" / "chosen-skill"
    skill_dir.mkdir(parents=True)
    (skill_dir / "SKILL.md").write_text("# Skill", encoding="utf-8")
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    assert server._find_compiled_skill("chosen-skill") == skill_dir.resolve()
    assert server._find_compiled_skill("../chosen-skill") is None
    assert server._find_compiled_skill("missing") is None


def test_upload_knowledge_writes_utf8_and_preserves_existing_file(tmp_path, monkeypatch):
    input_dir = tmp_path / "data" / "input"
    input_dir.mkdir(parents=True)
    (input_dir / "guide.md").write_text("existing", encoding="utf-8")
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    uploaded = "新的领域知识".encode("gb18030")
    result = server.save_uploaded_knowledge({
        "source_path": "data/input",
        "files": [{
            "name": "guide.md",
            "content_b64": base64.b64encode(uploaded).decode("ascii"),
        }],
    })

    assert result["ok"] is True
    assert result["written"] == [{
        "name": "guide-2.md",
        "path": "data/input/guide-2.md",
        "original_name": "guide.md",
        "size_bytes": len(uploaded),
        "normalized_size_bytes": len("新的领域知识\n".encode("utf-8")),
        "renamed": True,
        "converted": False,
        "source_format": "md",
        "source_encoding": "gb18030",
    }]
    assert (input_dir / "guide.md").read_text(encoding="utf-8") == "existing"
    assert (input_dir / "guide-2.md").read_text(encoding="utf-8") == "新的领域知识\n"
    assert (tmp_path / ".knowledge_originals" / "input" / "guide.md").read_bytes() == uploaded
    assert result["source"]["document_count"] == 2


def test_upload_knowledge_rejects_unsafe_or_binary_inputs(tmp_path, monkeypatch):
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    payload = base64.b64encode(b"safe text").decode("ascii")

    for source_path, filename in [
        ("../outside", "guide.md"),
        ("data/input", "../guide.md"),
        ("data/input", "guide.csv"),
        ("data/input", "guide.pdf"),
    ]:
        try:
            server.save_uploaded_knowledge({
                "source_path": source_path,
                "files": [{"name": filename, "content_b64": payload}],
            })
        except ValueError:
            pass
        else:
            raise AssertionError(f"unsafe upload was accepted: {source_path} / {filename}")


def test_upload_knowledge_converts_docx_to_markdown(tmp_path, monkeypatch):
    from docx import Document

    document = Document()
    document.add_heading("退款规则", level=1)
    document.add_paragraph("签收后七天内可以申请退款。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "场景"
    table.cell(0, 1).text = "时限"
    table.cell(1, 0).text = "退款"
    table.cell(1, 1).text = "7 天"
    payload = io.BytesIO()
    document.save(payload)

    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    result = server.save_uploaded_knowledge({
        "source_path": "data/input",
        "create_source": True,
        "files": [{
            "name": "售后手册.docx",
            "content_b64": base64.b64encode(payload.getvalue()).decode("ascii"),
        }],
    })

    markdown = (tmp_path / "data" / "input" / "售后手册.md").read_text(encoding="utf-8")
    assert markdown.startswith("# 售后手册\n")
    assert "## 退款规则" in markdown
    assert "签收后七天内可以申请退款。" in markdown
    assert "| 场景 | 时限 |" in markdown
    assert "| 退款 | 7 天 |" in markdown
    assert result["written"][0]["converted"] is True
    assert result["written"][0]["source_format"] == "docx"
    assert (
        tmp_path / ".knowledge_originals" / "input" / "售后手册.docx"
    ).read_bytes() == payload.getvalue()


def test_upload_knowledge_converts_xlsx_sheets_to_markdown(tmp_path, monkeypatch):
    from openpyxl import Workbook

    workbook = Workbook()
    rules = workbook.active
    rules.title = "规则"
    rules.append(["场景", "时限"])
    rules.append(["退款", 7])
    formulas = workbook.create_sheet("计算")
    formulas.append(["项目", "值"])
    formulas.append(["合计", "=SUM(1,2)"])
    payload = io.BytesIO()
    workbook.save(payload)
    workbook.close()

    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    result = server.save_uploaded_knowledge({
        "source_path": "data/input",
        "create_source": True,
        "files": [{
            "name": "规则表.xlsx",
            "content_b64": base64.b64encode(payload.getvalue()).decode("ascii"),
        }],
    })

    markdown = (tmp_path / "data" / "input" / "规则表.md").read_text(encoding="utf-8")
    assert markdown.startswith("# 规则表\n")
    assert "## 工作表：规则" in markdown
    assert "| 退款 | 7 |" in markdown
    assert "## 工作表：计算" in markdown
    assert "| 合计 | =SUM(1,2) |" in markdown
    assert result["written"][0]["converted"] is True
    assert result["written"][0]["source_format"] == "xlsx"
    assert (
        tmp_path / ".knowledge_originals" / "input" / "规则表.xlsx"
    ).read_bytes() == payload.getvalue()


def test_upload_knowledge_converts_text_pdf_and_rejects_scanned_pdf(tmp_path, monkeypatch):
    from pypdf import PdfWriter
    from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=200)
    font = DictionaryObject({
        NameObject("/Type"): NameObject("/Font"),
        NameObject("/Subtype"): NameObject("/Type1"),
        NameObject("/BaseFont"): NameObject("/Helvetica"),
    })
    resources = DictionaryObject({
        NameObject("/Font"): DictionaryObject({NameObject("/F1"): writer._add_object(font)}),
    })
    page[NameObject("/Resources")] = resources
    content = DecodedStreamObject()
    content.set_data(b"BT /F1 12 Tf 20 100 Td (Refund within 7 days) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(content)
    text_pdf = io.BytesIO()
    writer.write(text_pdf)

    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    result = server.save_uploaded_knowledge({
        "source_path": "data/input",
        "create_source": True,
        "files": [{
            "name": "policy.pdf",
            "content_b64": base64.b64encode(text_pdf.getvalue()).decode("ascii"),
        }],
    })

    markdown = (tmp_path / "data" / "input" / "policy.md").read_text(encoding="utf-8")
    assert "## 第 1 页" in markdown
    assert "Refund within 7 days" in markdown
    assert result["written"][0]["source_format"] == "pdf"

    scanned_writer = PdfWriter()
    scanned_writer.add_blank_page(width=300, height=200)
    scanned_pdf = io.BytesIO()
    scanned_writer.write(scanned_pdf)
    with pytest.raises(ValueError, match="OCR"):
        server.save_uploaded_knowledge({
            "source_path": "data/scanned",
            "create_source": True,
            "files": [{
                "name": "scan.pdf",
                "content_b64": base64.b64encode(scanned_pdf.getvalue()).decode("ascii"),
            }],
        })
    assert not (tmp_path / "data" / "scanned").exists()
    assert not (tmp_path / ".knowledge_originals" / "scanned").exists()


def test_empty_office_documents_are_rejected(tmp_path, monkeypatch):
    from docx import Document
    from openpyxl import Workbook

    empty_docx = io.BytesIO()
    Document().save(empty_docx)
    empty_xlsx = io.BytesIO()
    workbook = Workbook()
    workbook.save(empty_xlsx)
    workbook.close()

    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    for filename, raw in (
        ("empty.docx", empty_docx.getvalue()),
        ("empty.xlsx", empty_xlsx.getvalue()),
        ("empty.txt", b" \n\t"),
    ):
        with pytest.raises(ValueError, match="没有可用文本"):
            server.save_uploaded_knowledge({
                "source_path": "data/input",
                "create_source": True,
                "files": [{
                    "name": filename,
                    "content_b64": base64.b64encode(raw).decode("ascii"),
                }],
            })
    assert not (tmp_path / "data" / "input").exists()
    assert not (tmp_path / ".knowledge_originals" / "input").exists()


def test_new_source_upload_is_atomic_and_rejects_duplicate(tmp_path, monkeypatch):
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    payload = {
        "source_path": "data/new-source",
        "create_source": True,
        "batch_id": "batch-new-source",
        "files": [
            {
                "name": "one.txt",
                "content_b64": base64.b64encode(b"first").decode("ascii"),
            },
            {
                "name": "two.txt",
                "content_b64": base64.b64encode(b"second").decode("ascii"),
            },
        ],
    }

    result = server.save_uploaded_knowledge(payload)

    assert result["source"]["path"] == "data/new-source"
    assert result["source"]["document_count"] == 2
    assert (tmp_path / "data" / "new-source" / "one.md").is_file()
    assert (tmp_path / "data" / "new-source" / "two.md").is_file()
    with pytest.raises(ValueError, match="知识源已存在"):
        server.save_uploaded_knowledge(payload)
    assert sorted(path.name for path in (tmp_path / "data" / "new-source").glob("*.md")) == [
        "one.md",
        "two.md",
    ]


def test_failed_new_source_write_rolls_back_directory_and_archive(tmp_path, monkeypatch):
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    original_write = server._write_bytes_atomic
    write_count = 0

    def fail_markdown_write(target, payload):
        nonlocal write_count
        write_count += 1
        if write_count == 2:
            raise OSError("simulated disk failure")
        original_write(target, payload)

    monkeypatch.setattr(server, "_write_bytes_atomic", fail_markdown_write)

    with pytest.raises(OSError, match="simulated disk failure"):
        server.save_uploaded_knowledge({
            "source_path": "data/rollback-source",
            "create_source": True,
            "files": [{
                "name": "knowledge.txt",
                "content_b64": base64.b64encode(b"knowledge").decode("ascii"),
            }],
        })

    assert not (tmp_path / "data" / "rollback-source").exists()
    assert not (tmp_path / ".knowledge_originals" / "rollback-source").exists()


def test_source_is_unavailable_while_postprocessing_and_concurrent_upload_is_rejected(
    tmp_path,
    monkeypatch,
):
    source = tmp_path / "data" / "input"
    source.mkdir(parents=True)
    (source / "existing.md").write_text("existing", encoding="utf-8")
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    original_normalize = server.ki.normalize_knowledge_document
    conversion_started = threading.Event()
    allow_conversion = threading.Event()

    def slow_normalize(filename, raw):
        conversion_started.set()
        assert allow_conversion.wait(timeout=5)
        return original_normalize(filename, raw)

    monkeypatch.setattr(server.ki, "normalize_knowledge_document", slow_normalize)
    payload = {
        "source_path": "data/input",
        "batch_id": "batch-slow",
        "files": [{
            "name": "new.txt",
            "content_b64": base64.b64encode(b"new knowledge").decode("ascii"),
        }],
    }
    outcome = {}

    def upload():
        try:
            outcome["result"] = server.save_uploaded_knowledge(payload)
        except Exception as exc:  # pragma: no cover - asserted below
            outcome["error"] = exc

    worker = threading.Thread(target=upload)
    worker.start()
    assert conversion_started.wait(timeout=5)

    detail = server._input_source_detail(source)
    assert detail["ready"] is False
    assert detail["ingestion"]["status"] == "processing"
    assert detail["ingestion"]["batch_id"] == "batch-slow"
    assert detail["ingestion"]["current_file"] == "new.txt"
    assert "owner_pid" not in detail["ingestion"]
    assert "pending_outputs" not in detail["ingestion"]

    manager = server.mj.MiningJobManager(tmp_path, max_parallel=1)
    with pytest.raises(server.mj.MiningJobError, match="后处理"):
        manager._source_dir("data/input")
    with pytest.raises(ValueError, match="正在处理"):
        server.save_uploaded_knowledge({**payload, "batch_id": "batch-conflict"})

    allow_conversion.set()
    worker.join(timeout=5)
    assert not worker.is_alive()
    assert "error" not in outcome
    assert outcome["result"]["source"]["ready"] is True
    assert outcome["result"]["source"]["ingestion"]["status"] == "ready"
    assert (source / "new.md").is_file()
    assert manager._source_dir("data/input") == source


def test_failed_postprocessing_is_visible_and_successful_retry_recovers_source(
    tmp_path,
    monkeypatch,
):
    from pypdf import PdfWriter

    source = tmp_path / "data" / "input"
    source.mkdir(parents=True)
    (source / "existing.md").write_text("existing", encoding="utf-8")
    scanned_writer = PdfWriter()
    scanned_writer.add_blank_page(width=300, height=200)
    scanned_pdf = io.BytesIO()
    scanned_writer.write(scanned_pdf)
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    with pytest.raises(ValueError, match="OCR"):
        server.save_uploaded_knowledge({
            "source_path": "data/input",
            "batch_id": "batch-failed",
            "files": [{
                "name": "scan.pdf",
                "content_b64": base64.b64encode(scanned_pdf.getvalue()).decode("ascii"),
            }],
        })

    failed = server._input_source_detail(source)
    assert failed["ready"] is False
    assert failed["ingestion"]["status"] == "failed"
    assert "OCR" in failed["ingestion"]["error"]
    with pytest.raises(server.mj.MiningJobError, match="后处理失败"):
        server.mj.MiningJobManager(tmp_path)._source_dir("data/input")

    recovered = server.save_uploaded_knowledge({
        "source_path": "data/input",
        "batch_id": "batch-retry",
        "files": [{
            "name": "retry.txt",
            "content_b64": base64.b64encode(b"recovered").decode("ascii"),
        }],
    })
    assert recovered["source"]["ready"] is True
    assert recovered["source"]["ingestion"]["status"] == "ready"
    assert recovered["source"]["ingestion"]["batch_id"] == "batch-retry"


def test_interrupted_postprocessing_is_marked_failed_on_startup(tmp_path):
    source = tmp_path / "data" / "input"
    source.mkdir(parents=True)
    (source / "existing.md").write_text("existing", encoding="utf-8")
    partial = source / "partial.md"
    partial.write_text("partial batch", encoding="utf-8")
    preserved = source / "changed-after-marker.md"
    preserved.write_text("newer content", encoding="utf-8")
    server.ki.write_ingestion_state(
        tmp_path,
        "input",
        batch_id="batch-interrupted",
        status="processing",
        stage="writing",
        progress=80,
        pending_outputs=[
            {
                "path": "data/input/partial.md",
                "sha256": hashlib.sha256(b"partial batch").hexdigest(),
            },
            {
                "path": "data/input/changed-after-marker.md",
                "sha256": hashlib.sha256(b"older content").hexdigest(),
            },
        ],
    )

    assert server.ki.mark_interrupted_ingestions(tmp_path) == 1
    state = server.ki.read_ingestion_state(tmp_path, "input", has_documents=True)
    assert state["status"] == "failed"
    assert "中断" in state["error"]
    assert not partial.exists()
    assert preserved.read_text(encoding="utf-8") == "newer content"
    assert state["pending_outputs"] == []


def test_knowledge_sources_can_be_created_merged_and_deleted(tmp_path, monkeypatch):
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    created = server.create_knowledge_source({"name": "abc"})
    assert created["source"]["path"] == "data/abc"
    assert created["source"]["document_count"] == 0
    deleted_archive = tmp_path / ".knowledge_originals" / "abc"
    deleted_archive.mkdir(parents=True)
    (deleted_archive / "raw.docx").write_bytes(b"raw")

    source_a = tmp_path / "data" / "source-a"
    source_b = tmp_path / "data" / "source-b"
    source_a.mkdir(parents=True)
    source_b.mkdir(parents=True)
    (source_a / "guide.md").write_text("A", encoding="utf-8")
    (source_b / "guide.md").write_text("B", encoding="utf-8")
    (source_b / "nested").mkdir()
    (source_b / "nested" / "faq.txt").write_text("FAQ", encoding="utf-8")
    originals_a = tmp_path / ".knowledge_originals" / "source-a"
    originals_b = tmp_path / ".knowledge_originals" / "source-b"
    originals_a.mkdir(parents=True)
    originals_b.mkdir(parents=True)
    (originals_a / "guide.docx").write_bytes(b"docx-a")
    (originals_b / "guide.docx").write_bytes(b"docx-b")

    merged = server.merge_knowledge_sources({
        "source_paths": ["data/source-a", "data/source-b"],
        "target_name": "combined",
    })

    assert merged["source"]["path"] == "data/combined"
    assert merged["source"]["document_count"] == 3
    assert (tmp_path / "data" / "combined" / "guide.md").read_text(encoding="utf-8") == "A"
    assert (tmp_path / "data" / "combined" / "guide-2.md").read_text(encoding="utf-8") == "B"
    assert (tmp_path / "data" / "combined" / "nested" / "faq.txt").is_file()
    assert (
        tmp_path / ".knowledge_originals" / "combined" / "source-a" / "guide.docx"
    ).read_bytes() == b"docx-a"
    assert (
        tmp_path / ".knowledge_originals" / "combined" / "source-b" / "guide.docx"
    ).read_bytes() == b"docx-b"
    assert (source_a / "guide.md").is_file()
    assert (source_b / "guide.md").is_file()

    deleted = server.delete_knowledge_source("abc")
    assert deleted["deleted"]["path"] == "data/abc"
    assert not (tmp_path / "data" / "abc").exists()
    assert not deleted_archive.exists()


def test_knowledge_source_can_be_renamed_without_overwriting_duplicate(tmp_path, monkeypatch):
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)
    original = tmp_path / "data" / "support-docs"
    original.mkdir(parents=True)
    (original / "guide.md").write_text("keep me", encoding="utf-8")
    original_archive = tmp_path / ".knowledge_originals" / "support-docs"
    original_archive.mkdir(parents=True)
    (original_archive / "guide.docx").write_bytes(b"original")
    existing = tmp_path / "data" / "existing"
    existing.mkdir()

    renamed = server.rename_knowledge_source("support-docs", {"name": "customer-support"})

    assert renamed["previous_path"] == "data/support-docs"
    assert renamed["source"]["path"] == "data/customer-support"
    assert not original.exists()
    assert (tmp_path / "data" / "customer-support" / "guide.md").read_text(encoding="utf-8") == "keep me"
    assert not original_archive.exists()
    assert (
        tmp_path / ".knowledge_originals" / "customer-support" / "guide.docx"
    ).read_bytes() == b"original"

    with pytest.raises(FileExistsError, match="知识源名称已存在"):
        server.rename_knowledge_source("customer-support", {"name": "EXISTING"})

    assert (tmp_path / "data" / "customer-support" / "guide.md").is_file()
    assert existing.is_dir()


def test_job_artifact_reader_does_not_expose_snapshot_or_metadata(tmp_path, monkeypatch):
    artifact = tmp_path / "mining_jobs" / "job-1" / "workspace" / "compiled_skill" / "demo" / "SKILL.md"
    artifact.parent.mkdir(parents=True)
    artifact.write_text("# Skill", encoding="utf-8")
    metadata = tmp_path / "mining_jobs" / "job-1" / "job.json"
    metadata.write_text('{"secret": "not-an-artifact"}', encoding="utf-8")
    snapshot = tmp_path / "mining_jobs" / "job-1" / "workspace" / "data" / "input" / "private.md"
    snapshot.parent.mkdir(parents=True)
    snapshot.write_text("private source", encoding="utf-8")
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    assert server.read_history_artifact(
        "mining_jobs/job-1/workspace/compiled_skill/demo/SKILL.md"
    )["content"] == "# Skill"
    for forbidden in (
        "mining_jobs/job-1/job.json",
        "mining_jobs/job-1/workspace/data/input/private.md",
    ):
        try:
            server.read_history_artifact(forbidden)
        except ValueError:
            pass
        else:
            raise AssertionError(f"non-artifact job file was exposed: {forbidden}")


def test_completed_job_artifact_can_be_human_edited(tmp_path, monkeypatch):
    job_dir = tmp_path / "mining_jobs" / "job-1"
    artifact = job_dir / "workspace" / "compiled_skill" / "demo" / "SKILL.md"
    artifact.parent.mkdir(parents=True)
    artifact.write_text("# Before", encoding="utf-8")
    (job_dir / "job.json").write_text(
        json.dumps({"job_id": "job-1", "status": "succeeded"}),
        encoding="utf-8",
    )
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    result = server.save_history_artifact(
        "mining_jobs/job-1/workspace/compiled_skill/demo/SKILL.md",
        "# After\n\nHuman refined content.\n",
    )

    assert result["edited"] is True
    assert result["content"].startswith("# After")
    assert artifact.read_text(encoding="utf-8") == result["content"]


def test_active_job_artifact_cannot_be_edited(tmp_path, monkeypatch):
    job_dir = tmp_path / "mining_jobs" / "job-1"
    artifact = job_dir / "workspace" / "compiled_skill" / "demo" / "SKILL.md"
    artifact.parent.mkdir(parents=True)
    artifact.write_text("# Original", encoding="utf-8")
    (job_dir / "job.json").write_text(
        json.dumps({"job_id": "job-1", "status": "running"}),
        encoding="utf-8",
    )
    monkeypatch.setattr(server, "PROJECT_ROOT", tmp_path)

    with pytest.raises(ValueError, match="已完成"):
        server.save_history_artifact(
            "mining_jobs/job-1/workspace/compiled_skill/demo/SKILL.md",
            "# Unauthorized edit",
        )
    assert artifact.read_text(encoding="utf-8") == "# Original"


def test_gap_questions_are_concrete_form_items(tmp_path):
    skill = tmp_path / "SKILL.md"
    skill.write_text(
        "## 维度一：时效判定\n"
        "- **高严重度缺口**：升级处理的时效指标缺失——请问超过多少小时必须升级？（来源：SOP 第 3 节）\n"
        "- **中严重度缺口**：特殊订单的判定标准未明确（来源：规则表）\n",
        encoding="utf-8",
    )

    questions = server.extract_gap_questions_from_skill(skill)

    assert len(questions) == 2
    assert questions[0]["question"] == "请问超过多少小时必须升级？"
    assert questions[0]["answer_type"] == "short_text"
    assert "单位" in questions[0]["field_label"]
    assert questions[0]["context"].startswith("升级处理的时效指标缺失")
    assert questions[0]["source"] == "SOP 第 3 节"
    assert questions[1]["question"].startswith("请明确“特殊订单的判定标准”")
    assert questions[1]["answer_type"] == "long_text"


def test_skillminer_ui_omits_redundant_microcopy():
    repository_root = Path(__file__).resolve().parents[1]
    mining_view = (repository_root / "web-ui" / "src" / "views" / "MiningView.tsx").read_text(
        encoding="utf-8"
    )
    legacy_console = (
        repository_root / "teamEvolver" / "skillminer" / "web_console" / "static" / "index.html"
    ).read_text(encoding="utf-8")

    for redundant in (
        "候选输入",
        "点击名称直接修改",
        "点击文件可预览并人工修订",
        "选择知识源和反思轮数",
        "创建后可在当前任务条目下查看实时进度与日志",
        "将创建知识源",
        "可以从不同知识源连续创建多个任务",
        "任务启动时固化输入，后续上传或合并不会影响本次挖掘",
        "可以连续新建多个任务，调度器自动运行或排队",
        "每个任务的 Skill 与 Benchmark 独立保存并可追溯",
        "独立任务 · 可并行",
        "模型执行组件由系统内置维护",
    ):
        assert redundant not in mining_view
    assert "对已编译的 skill 做 benchmark 跑分" not in legacy_console

    # Non-obvious safeguards and destructive-action consequences remain visible.
    assert "不会覆盖同名文件" in mining_view
    assert "个文档会一并删除" in mining_view
    assert "超出后自动排队" in mining_view


def test_new_source_upload_snapshots_files_and_uses_one_transaction():
    repository_root = Path(__file__).resolve().parents[1]
    mining_view = (repository_root / "web-ui" / "src" / "views" / "MiningView.tsx").read_text(
        encoding="utf-8"
    )
    upload_handler = mining_view.split(
        "async function uploadKnowledgeFiles", 1
    )[1].split("async function startRun", 1)[0]

    assert upload_handler.index("const selectedFiles = Array.from(list)") < upload_handler.index(
        "await Promise.all(selectedFiles.map"
    )
    assert 'api<{ source: InputSource }>("/api/mining/sources"' not in upload_handler
    assert "create_source: createSource" in upload_handler


def test_mining_overview_exposes_source_totals_and_parallel_job_pipelines():
    repository_root = Path(__file__).resolve().parents[1]
    mining_view = (repository_root / "web-ui" / "src" / "views" / "MiningView.tsx").read_text(
        encoding="utf-8"
    )
    overview = mining_view.split('{/* ---- 总览 ---- */}', 1)[1].split(
        '{/* ---- 知识源 ---- */}', 1
    )[0]

    assert 'label="知识源"' in overview
    assert 'label="文档总数"' in overview
    assert 'title="知识源文档"' not in overview
    assert "source.document_count" not in overview
    assert 'title="任务流水线"' in overview
    assert "pipelineJobs.map" in overview
    assert "job.phase?.[step.key]" in overview
    assert '<Panel title="近期任务">' in overview
    assert "总任务 ${jobCounts.total} · 展示 ${recentTaskCount}" not in overview
    assert "mining.jobs.slice(0, recentTaskCount)" in overview


def test_skillminer_user_facing_copy_uses_knowledge_source_terminology():
    repository_root = Path(__file__).resolve().parents[1]
    legacy_term = "数" + "据源"
    for relative_path in (
        "web-ui/src/views/MiningView.tsx",
        "teamEvolver/skillminer/web_console/server.py",
        "teamEvolver/skillminer/run_coverage_report.py",
    ):
        contents = (repository_root / relative_path).read_text(encoding="utf-8")
        assert legacy_term not in contents
